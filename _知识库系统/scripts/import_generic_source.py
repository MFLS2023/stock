#!/usr/bin/env python3
"""Generic importer for registered MD/TXT/PDF/DOCX/image source folders."""

from __future__ import annotations

import argparse
import hashlib
import json
import shutil
import subprocess
import tempfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

import yaml
from docx import Document
from pypdf import PdfReader

from kb_import_utils import (
    ROOT,
    clean_text,
    clean_title,
    extract_date,
    infer_topics,
    meaningful_char_count,
    natural_key,
    ocr_images,
    split_text,
    write_jsonl,
    write_text_lf,
)


CONFIG = ROOT / "_知识库系统" / "config" / "sources.yaml"
POPPLER = Path(
    r"C:\Users\20577\.cache\codex-runtimes\codex-primary-runtime\dependencies\native\poppler\Library\bin\pdftoppm.exe"
)
SUPPORTED = {".md", ".txt", ".pdf", ".docx", ".jpg", ".jpeg", ".png"}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def render_page(pdf: Path, page: int, output: Path, dpi: int) -> None:
    executable = str(POPPLER) if POPPLER.exists() else (shutil.which("pdftoppm") or "")
    if not executable:
        raise FileNotFoundError("pdftoppm not found")
    output.parent.mkdir(parents=True, exist_ok=True)
    completed = subprocess.run(
        [executable, "-png", "-r", str(dpi), "-f", str(page), "-l", str(page), "-singlefile", str(pdf), str(output.with_suffix(""))],
        cwd=str(Path(executable).parent), capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    if completed.returncode != 0 or not output.exists():
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "PDF render failed")


def extract_pdf(path: Path, doc_id: str, digest: str, lib: Path, force: bool, dpi: int, threshold: int):
    reader = PdfReader(path)
    units: list[dict] = []
    errors: list[dict] = []
    work = Path(tempfile.mkdtemp(prefix=f"generic-{doc_id}-", dir=ROOT / "_知识库系统" / "tmp"))
    pending: list[tuple[str, Path]] = []
    page_meta: dict[int, dict] = {}
    try:
        for index, page in enumerate(reader.pages, start=1):
            cache = lib / "page_texts" / f"{doc_id}-page-{index:03d}.json"
            if not force and cache.exists():
                try:
                    item = json.loads(cache.read_text(encoding="utf-8"))
                    if item.get("source_hash") == digest and item.get("dpi") == dpi:
                        page_meta[index] = item
                        continue
                except (OSError, json.JSONDecodeError):
                    pass
            try:
                embedded = clean_text(page.extract_text() or "")
            except Exception as exc:
                embedded = ""
                errors.append({"page": index, "stage": "embedded", "error": str(exc)})
            if meaningful_char_count(embedded) >= threshold:
                page_meta[index] = {"source_hash": digest, "dpi": dpi, "text": embedded, "method": "embedded", "confidence": "high"}
            else:
                try:
                    image = work / f"page-{index:03d}.png"
                    render_page(path, index, image, dpi)
                    pending.append((str(index), image))
                    page_meta[index] = {"source_hash": digest, "dpi": dpi, "text": embedded, "method": "pending", "confidence": "low"}
                except Exception as exc:
                    errors.append({"page": index, "stage": "render", "error": str(exc)})
                    page_meta[index] = {"source_hash": digest, "dpi": dpi, "text": embedded, "method": "embedded_fallback", "confidence": "low"}
        if pending:
            results = ocr_images(pending)
            for key, _ in pending:
                index = int(key)
                embedded = page_meta[index]["text"]
                result = results.get(key, {"text": "", "error": "no_result"})
                ocr_text = clean_text(result.get("text", ""), ocr=True)
                use_ocr = meaningful_char_count(ocr_text) >= max(80, int(meaningful_char_count(embedded) * 0.75))
                page_meta[index] = {
                    "source_hash": digest, "dpi": dpi, "text": ocr_text if use_ocr else embedded,
                    "method": "ocr" if use_ocr else "embedded_fallback", "confidence": "medium" if use_ocr else "low",
                    "ocr_error": result.get("error", ""),
                }
                if result.get("error"):
                    errors.append({"page": index, "stage": "ocr", "error": result["error"]})
        for index in range(1, len(reader.pages) + 1):
            item = page_meta[index]
            cache = lib / "page_texts" / f"{doc_id}-page-{index:03d}.json"
            write_text_lf(cache, json.dumps(item, ensure_ascii=False, indent=2) + "\n")
            if item.get("text"):
                units.append({"locator": f"第{index}页", "text": item["text"], "method": item["method"], "confidence": item["confidence"]})
    finally:
        shutil.rmtree(work, ignore_errors=True)
    return units, errors


def extract_docx(path: Path):
    document = Document(path)
    units = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = clean_text(paragraph.text)
        if text:
            units.append({"locator": f"正文第{index}段", "text": text, "method": "docx", "confidence": "high"})
    for table_index, table in enumerate(document.tables, start=1):
        rows = [" | ".join(clean_text(cell.text) for cell in row.cells) for row in table.rows]
        text = clean_text("\n".join(rows))
        if text:
            units.append({"locator": f"表格{table_index}", "text": text, "method": "docx_table", "confidence": "high"})
    for section_index, section in enumerate(document.sections, start=1):
        for label, paragraphs in (("页眉", section.header.paragraphs), ("页脚", section.footer.paragraphs)):
            text = clean_text("\n".join(paragraph.text for paragraph in paragraphs))
            if text:
                units.append({"locator": f"第{section_index}节{label}", "text": text, "method": "docx", "confidence": "high"})
    return units, len(document.inline_shapes)


def group_units(units: list[dict], target: int):
    groups, current, length = [], [], 0
    for unit in units:
        if current and length + len(unit["text"]) > target:
            groups.append(current)
            current, length = [], 0
        current.append(unit)
        length += len(unit["text"])
    if current:
        groups.append(current)
    return groups


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", required=True)
    parser.add_argument("--force", action="store_true")
    parser.add_argument("--dpi", type=int, default=140)
    parser.add_argument("--embedded-threshold", type=int, default=500)
    args = parser.parse_args()
    config = yaml.safe_load(CONFIG.read_text(encoding="utf-8"))
    source = next((item for item in config.get("sources", []) if item["id"] == args.source), None)
    if source is None:
        raise KeyError(args.source)
    source_root = ROOT / source["source_path"]
    lib = ROOT / "_知识库系统" / "source_libraries" / source["id"]
    for name in ("texts", "page_texts", "image_ocr_cache", "maps"):
        (lib / name).mkdir(parents=True, exist_ok=True)
    files = sorted((path for path in source_root.rglob("*") if path.is_file()), key=lambda path: natural_key(path.relative_to(source_root).as_posix()))
    supported_files = [path for path in files if path.suffix.lower() in SUPPORTED]
    unsupported = [path for path in files if path.suffix.lower() not in SUPPORTED]
    hashes = {path: sha256_file(path) for path in supported_files}
    canonical = {}
    for path in supported_files:
        canonical.setdefault(hashes[path], path)

    documents, parents, chunks, errors = [], [], [], []
    used_ids: Counter[str] = Counter()
    extraction_counts: Counter[str] = Counter()
    for path in supported_files:
        digest = hashes[path]
        base_id = f"{source['id']}-{digest[:12]}"
        used_ids[base_id] += 1
        doc_id = base_id if used_ids[base_id] == 1 else f"{base_id}-dup{used_ids[base_id]}"
        duplicate_of = base_id if path != canonical[digest] else ""
        units = []
        embedded_images = 0
        if not duplicate_of:
            try:
                suffix = path.suffix.lower()
                if suffix in {".md", ".txt"}:
                    units = [{"locator": "全文", "text": clean_text(read_text_file(path)), "method": suffix[1:], "confidence": "high"}]
                elif suffix == ".pdf":
                    units, page_errors = extract_pdf(path, doc_id, digest, lib, args.force, args.dpi, args.embedded_threshold)
                    errors.extend({"document_id": doc_id, **item} for item in page_errors)
                elif suffix == ".docx":
                    units, embedded_images = extract_docx(path)
                else:
                    cache = lib / "image_ocr_cache" / f"{digest}.json"
                    if not args.force and cache.exists():
                        result = json.loads(cache.read_text(encoding="utf-8"))
                    else:
                        raw = ocr_images([("image", path)]).get("image", {"text": "", "error": "no_result"})
                        result = {"source_hash": digest, "text": clean_text(raw.get("text", ""), ocr=True), "error": raw.get("error", "")}
                        write_text_lf(cache, json.dumps(result, ensure_ascii=False, indent=2) + "\n")
                    units = [{"locator": f"图片:{path.name}", "text": result.get("text", ""), "method": "ocr", "confidence": "medium" if result.get("text") else "low"}]
                    if result.get("error"):
                        errors.append({"document_id": doc_id, "stage": "ocr", "error": result["error"]})
            except Exception as exc:
                errors.append({"document_id": doc_id, "stage": "extract", "error": str(exc)})
                units = []
        expanded = []
        for unit in units:
            if not unit.get("text"):
                continue
            extraction_counts[unit["method"]] += 1
            expanded.extend({**unit, "text": part} for part in split_text(unit["text"], 1200))
        normalized = "\n\n".join(f"[{unit['locator']} | {unit['method']}]\n{unit['text']}" for unit in expanded)
        text_path = lib / "texts" / f"{doc_id}.txt"
        write_text_lf(text_path, normalized + ("\n" if normalized else ""))
        title = clean_title(path)
        date = extract_date(path.name, path.relative_to(source_root).as_posix(), normalized[:500])
        topics = infer_topics(title, normalized)
        risk = []
        if any(unit["method"] == "ocr" for unit in expanded):
            risk.append("OCR内容需回看原文件")
        if embedded_images:
            risk.append(f"DOCX含{embedded_images}个内嵌图，通用适配器未OCR，建议专用适配器")
        documents.append({
            "source_id": source["id"], "document_id": doc_id, "title": title, "date": date,
            "author_or_guest": source.get("author") or source["display_name"], "content_type": path.suffix.lower().lstrip("."),
            "topics": topics, "characters": meaningful_char_count(normalized), "original_path": str(path),
            "normalized_text_path": str(text_path), "sha256": digest, "duplicate_of": duplicate_of,
            "risk_flags": "；".join(risk),
        })
        if duplicate_of:
            continue
        for parent_number, group in enumerate(group_units(expanded, 4200), start=1):
            parent_id = f"{doc_id}-p{parent_number:03d}"
            locator = group[0]["locator"] if len(group) == 1 else f"{group[0]['locator']}—{group[-1]['locator']}"
            parents.append({
                "source_id": source["id"], "document_id": doc_id, "parent_id": parent_id, "title": title,
                "date": date, "author_or_guest": source.get("author") or source["display_name"],
                "locator": locator, "text": "\n".join(item["text"] for item in group),
            })
            for child_number, unit in enumerate(group, start=1):
                chunks.append({
                    "source_id": source["id"], "source_name": source["display_name"], "document_id": doc_id,
                    "parent_id": parent_id, "chunk_id": f"{parent_id}-c{child_number:02d}", "chunk_type": "generic_document",
                    "title": title, "date": date, "author_or_guest": source.get("author") or source["display_name"],
                    "topics": topics, "claim_type": "opinion_or_case", "market_regime": "未标注",
                    "locator": unit["locator"], "text": unit["text"], "original_path": str(path),
                    "confidence": unit["confidence"], "extraction_method": unit["method"],
                })

    write_jsonl(lib / "documents.jsonl", documents)
    write_jsonl(lib / "parents.jsonl", parents)
    write_jsonl(lib / "chunks.jsonl", chunks)
    content_map = [f"# {source['display_name']} 内容地图", "", "| 日期 | 标题 | 格式 | 主题 | 字符数 |", "|---|---|---|---|---:|"]
    for item in sorted(documents, key=lambda row: (row.get("date") or "9999", row["title"])):
        content_map.append(f"| {item.get('date') or '未标注'} | {item['title'].replace('|', '｜')} | {item['content_type']} | {'、'.join(item.get('topics') or []) or '未标注'} | {item['characters']} |")
    write_text_lf(lib / "content_map.md", "\n".join(content_map) + "\n")
    summary = {
        "generated_at": datetime.now(timezone.utc).isoformat(), "source_id": source["id"], "physical_files": len(files),
        "supported_files": len(supported_files), "unsupported_files": [str(path) for path in unsupported],
        "documents": len(documents), "parents": len(parents), "chunks": len(chunks),
        "extraction_counts": dict(extraction_counts), "errors": errors,
    }
    write_text_lf(lib / "source_summary.json", json.dumps(summary, ensure_ascii=False, indent=2) + "\n")
    quality = [
        f"# {source['display_name']} 导入质量报告", "", f"- 物理文件：{len(files)}", f"- 已支持：{len(supported_files)}",
        f"- 暂不支持：{len(unsupported)}", f"- 文档：{len(documents)}", f"- 检索块：{len(chunks)}",
        f"- 错误：{len(errors)}", "", "通用适配器适用于常规 MD/TXT/PDF/DOCX/图片。若存在说话人、复杂图文关系、跨图顺序或关键图表，应升级为专用适配器。",
    ]
    write_text_lf(lib / "quality_report.md", "\n".join(quality) + "\n")
    source_yaml = lib / "source.yaml"
    updated_source = dict(source)
    updated_source["status"] = "integrated" if not errors and not unsupported else "integrated_with_warnings"
    write_text_lf(source_yaml, yaml.safe_dump(updated_source, allow_unicode=True, sort_keys=False))
    source["status"] = updated_source["status"]
    source["last_import_summary"] = {
        "documents": len(documents), "chunks": len(chunks), "errors": len(errors),
        "unsupported": len(unsupported),
    }
    temporary_config = CONFIG.with_suffix(".yaml.tmp")
    write_text_lf(temporary_config, yaml.safe_dump(config, allow_unicode=True, sort_keys=False))
    temporary_config.replace(CONFIG)
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0 if not errors and not unsupported else 1


if __name__ == "__main__":
    raise SystemExit(main())
